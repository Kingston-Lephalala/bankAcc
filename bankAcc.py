from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Account
class BankAccount:
    def __init__(self, name, pin, bal=0.0):
        self.holder = name
        self.pin = pin
        self.balance = bal
        self.transactions = []  # List to store transaction history

    def deposit(self, amount):
        self.balance += amount
        transaction = [datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "Deposit", f"R{amount:.2f}", f"R{self.balance:.2f}"]
        self.transactions.append(transaction)
        print(f"Deposited: R{amount:.2f}. New balance is R{self.balance:.2f}\n")
        self.save_transactions_to_docx()  # Save to DOCX after each transaction

    def withdraw(self, amount):
        pin_attempt = input("Enter your PIN to withdraw: ")
        if pin_attempt != self.pin:
            print("Incorrect PIN! Withdrawal denied.\n")
            return
        
        if amount > self.balance:
            print(f"Insufficient balance! You only have R{self.balance:.2f}\n")
        else:
            self.balance -= amount
            transaction = [datetime.now().strftime("%Y-%m-%d %H:%M:%S"), "Withdraw", f"R{amount:.2f}", f"R{self.balance:.2f}"]
            self.transactions.append(transaction)
            print(f"Withdrew amount: R{amount:.2f}. New Balance: R{self.balance:.2f}\n")
            self.save_transactions_to_docx()  # Save to DOCX after each transaction

    def display(self):
        pin_attempt = input("Enter your PIN to check balance: ")
        if pin_attempt != self.pin:
            print("Incorrect PIN! Access denied.\n")
            return
        print(f"\nCurrent balance: R{self.balance:.2f}\n")

    def show_transaction_history(self):
        pin_attempt = input("Enter your PIN to view transaction history: ")
        if pin_attempt != self.pin:
            print("Incorrect PIN! Access denied.\n")
            return
        
        print("\nTransaction History:")
        if not self.transactions:
            print("No transactions yet.\n")
        else:
            for t in self.transactions:
                print(f"{t[0]} | {t[1]} | {t[2]} | Balance: {t[3]}")
        print()

    def save_transactions_to_docx(self):
        doc = Document()

        # **Title Formatting**
        title = doc.add_paragraph()
        title_run = title.add_run(f"Transaction History for {self.holder}")
        title_run.bold = True
        title_run.font.size = Pt(18)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("\n")  # Add space after title

        # **Create Table**
        if not self.transactions:
            doc.add_paragraph("No transactions yet.")
        else:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"

            # **Table Header**
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Date & Time"
            hdr_cells[1].text = "Type"
            hdr_cells[2].text = "Amount"
            hdr_cells[3].text = "Balance"

            # **Formatting Header**
            for cell in hdr_cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # **Adding Transactions to Table**
            for transaction in self.transactions:
                row_cells = table.add_row().cells
                for i, text in enumerate(transaction):
                    row_cells[i].text = text

        # **Save the document**
        filename = f"{self.holder}_transaction_history.docx"
        doc.save(filename)
        print(f"Transaction history saved to {filename}\n")


# Open account
userName = input("Enter Name: ")
userPin = input("Set a 4-digit PIN: ")
defAmnt = float(input("Enter starting amount: R"))
account = BankAccount(userName, userPin, defAmnt)

while True:
    prompt = input("Enter D for Deposit\nEnter W for Withdrawal\nEnter B for Balance\nEnter H for History\nEnter X for exit\n").upper()

    if prompt == "D":
        depAmnt = float(input("\nHow much do you want to deposit: R"))
        account.deposit(depAmnt)

    elif prompt == "W":
        witAmnt = float(input("\nHow much do you want to withdraw: R"))
        account.withdraw(witAmnt)

    elif prompt == "B":
        account.display()

    elif prompt == "H":
        account.show_transaction_history()

    elif prompt == "X":
        print("\nExiting. Thank you for using our bank!\n")
        break

    else:
        print("Invalid option! Please enter D, W, B, H, or X.\n")
